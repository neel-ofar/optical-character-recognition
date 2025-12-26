
from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
import cv2
import numpy as np
import pytesseract
from PIL import Image
import os
from datetime import datetime
from werkzeug.utils import secure_filename
from pathlib import Path
import tempfile
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from googletrans import Translator
import pyttsx3
from pydub import AudioSegment
import re
import heapq
from collections import defaultdict

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
RESULTS_FOLDER = 'results'
EXPORT_FOLDER = 'exports'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'pdf'}
MAX_FILE_SIZE = 100 * 1024 * 1024

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['RESULTS_FOLDER'] = RESULTS_FOLDER
app.config['EXPORT_FOLDER'] = EXPORT_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

Path(UPLOAD_FOLDER).mkdir(exist_ok=True)
Path(RESULTS_FOLDER).mkdir(exist_ok=True)
Path(EXPORT_FOLDER).mkdir(exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def preprocess_image(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    denoised = cv2.fastNlMeansDenoising(gray)
    thresh = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
    return thresh

def extract_text_from_image(image_path, lang='eng'):
    try:
        image = cv2.imread(image_path)
        processed = preprocess_image(image)
        text = pytesseract.image_to_string(processed, lang=lang)
        data = pytesseract.image_to_data(processed, lang=lang, output_type=pytesseract.Output.DICT)
        confidences = [int(c) for c in data['conf'] if c != '-1']
        avg_confidence = np.mean(confidences) if confidences else 0
        return {'text': text.strip(), 'confidence': float(avg_confidence), 'success': True}
    except Exception as e:
        return {'text': '', 'confidence': 0, 'success': False, 'error': str(e)}

def extract_text_from_pdf(pdf_path, lang='eng'):
    try:
        all_text = []
        total_confidence = []
        images = convert_from_path(pdf_path, dpi=300)
        for i, image in enumerate(images):
            temp_image = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            image.save(temp_image.name, 'PNG')
            result = extract_text_from_image(temp_image.name, lang=lang)
            if result['success']:
                all_text.append(f"\n=== Page {i+1} ===\n{result['text']}")
                total_confidence.append(result['confidence'])
            os.unlink(temp_image.name)
        combined_text = '\n'.join(all_text)
        avg_confidence = np.mean(total_confidence) if total_confidence else 0
        return {'text': combined_text.strip(), 'confidence': float(avg_confidence), 'pages': len(images), 'success': True}
    except Exception as e:
        return {'text': '', 'confidence': 0, 'pages': 0, 'success': False, 'error': str(e)}

def translate_text(text, target_lang):
    try:
        translator = Translator()
        translated = translator.translate(text, dest=target_lang)
        return translated.text, True
    except Exception as e:
        return text, False

def create_word_document(text, filename, is_translated=False, translate_lang=None):
    doc = Document()
    title_text = 'Translated OCR Extracted Text' if is_translated else 'OCR Extracted Text'
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta = doc.add_paragraph()
    meta.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n').italic = True
    meta.add_run(f'Source: {filename}\n').italic = True
    if is_translated and translate_lang:
        meta.add_run(f'Translated to: {translate_lang}\n').italic = True
    doc.add_heading('Extracted Content', level=1)
    for para in text.split('\n\n'):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.style.font.size = Pt(11)
    suffix = '_translated' if is_translated else ''
    stem = Path(filename).stem
    output_path = os.path.join(EXPORT_FOLDER, f"{stem}{suffix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(output_path)
    return output_path

def create_text_file(text, filename, is_translated=False, translate_lang=None):
    suffix = '_translated' if is_translated else ''
    stem = Path(filename).stem
    output_path = os.path.join(EXPORT_FOLDER, f"{stem}{suffix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")
    with open(output_path, 'w', encoding='utf-8') as f:
        title_text = 'Translated OCR Extracted Text' if is_translated else 'OCR Extracted Text'
        f.write(f"{title_text}\n{'='*60}\n\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Source: {filename}\n")
        if is_translated and translate_lang:
            f.write(f"Translated to: {translate_lang}\n")
        f.write(f"\n{'='*60}\n\n")
        f.write(text)
    return output_path

def create_audio_file(text, gender, tone, filename):
    engine = pyttsx3.init()
    emotions = {
        'professional': {'rate': 150, 'volume': 1.0},
        'friendly':     {'rate': 160, 'volume': 1.0},
        'happy':        {'rate': 180, 'volume': 1.0},
        'excited':      {'rate': 200, 'volume': 1.2},
        'sad':          {'rate': 120, 'volume': 0.8},
        'enthusiastic': {'rate': 190, 'volume': 1.1},
        'romantic':     {'rate': 140, 'volume': 0.9},
        'neutral':      {'rate': 150, 'volume': 1.0},
    }
    params = emotions.get(tone.lower(), emotions['neutral'])
    engine.setProperty('rate', params['rate'])
    engine.setProperty('volume', params['volume'])
    voices = engine.getProperty('voices')
    selected_voice = next((v.id for v in voices if gender.lower() in v.name.lower()), voices[0].id if voices else None)
    if selected_voice:
        engine.setProperty('voice', selected_voice)
    wav_path = tempfile.NamedTemporaryFile(delete=False, suffix='.wav').name
    engine.save_to_file(text, wav_path)
    engine.runAndWait()
    audio = AudioSegment.from_wav(wav_path)
    stem = Path(filename).stem
    output_path = os.path.join(EXPORT_FOLDER, f"{stem}_audio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp3")
    audio.export(output_path, format='mp3')
    os.unlink(wav_path)
    return output_path

def summarize_text(text, mode='brief'):
    """Simple extractive summarization using sentence scoring"""
    if not text.strip():
        return "No text to summarize."

    # Clean and split into sentences
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    if len(sentences) <= 1:
        return text.strip()

    # Remove common stop words
    stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'been', 'be', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'what', 'which', 'who', 'when', 'where', 'why', 'how', 'all', 'each', 'few', 'more', 'most', 'some', 'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 's', 't', 'just', 'don', 'now'}

    # Word frequencies
    word_frequencies = defaultdict(int)
    for sentence in sentences:
        for word in re.findall(r'\w+', sentence.lower()):
            if word not in stop_words:
                word_frequencies[word] += 1

    # Score sentences
    sentence_scores = {}
    for sentence in sentences:
        for word in re.findall(r'\w+', sentence.lower()):
            if word in word_frequencies:
                if len(sentence.split()) < 30:  # avoid too long sentences
                    if sentence not in sentence_scores:
                        sentence_scores[sentence] = word_frequencies[word]
                    else:
                        sentence_scores[sentence] += word_frequencies[word]

    # Get top sentences
    if mode == 'brief':
        summary_sentences = heapq.nlargest(3, sentence_scores, key=sentence_scores.get)
    elif mode == 'detailed':
        summary_sentences = heapq.nlargest(min(8, len(sentences)//2), sentence_scores, key=sentence_scores.get)
    elif mode == 'bullet':
        summary_sentences = heapq.nlargest(7, sentence_scores, key=sentence_scores.get)
    else:
        summary_sentences = heapq.nlargest(4, sentence_scores, key=sentence_scores.get)

    # Preserve order
    summary_sentences = sorted(summary_sentences, key=lambda s: sentences.index(s))

    if mode == 'bullet':
        return '\n'.join(f"• {s.strip()}" for s in summary_sentences if s.strip())
    else:
        return ' '.join(summary_sentences)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/ocr', methods=['POST'])
def ocr_endpoint():
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file uploaded'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    if not allowed_file(file.filename):
        return jsonify({'success': False, 'error': 'Invalid file type'}), 400

    ocr_lang = request.form.get('ocr_lang', 'eng')
    translate_to = request.form.get('translate_to', '')

    try:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = secure_filename(file.filename)
        stem = Path(filename).stem
        ext = Path(filename).suffix.lower()
        saved_filename = f"{stem}_{timestamp}{ext}"
        filepath = os.path.join(UPLOAD_FOLDER, saved_filename)
        file.save(filepath)

        if ext == '.pdf':
            result = extract_text_from_pdf(filepath, lang=ocr_lang)
        else:
            result = extract_text_from_image(filepath, lang=ocr_lang)

        translated = False
        if result['success'] and translate_to:
            result['text'], translated = translate_text(result['text'], translate_to)
            result['translated'] = translated
            result['translate_lang'] = translate_to if translated else None
            if result['success']:
                text = result['text']
                word_count = len(text.split())
                char_count = len(text)  # includes spaces
                result['word_count'] = word_count
                result['char_count'] = char_count

        result['translated'] = translated
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/summarize', methods=['POST'])
def summarize_endpoint():
    data = request.get_json()
    text = data.get('text', '').strip()
    mode = data.get('mode', 'brief')  # brief, detailed, bullet

    if not text:
        return jsonify({'error': 'No text to summarize'}), 400

    summary = summarize_text(text, mode=mode)
    return jsonify({'summary': summary, 'mode': mode})

@app.route('/api/export/txt', methods=['POST'])
def export_txt():
    data = request.get_json()
    path = create_text_file(
        data.get('text', ''),
        data.get('filename', 'document'),
        data.get('translated', False),
        data.get('translate_lang')
    )
    return send_file(path, as_attachment=True, mimetype='text/plain')

@app.route('/api/export/docx', methods=['POST'])
def export_docx():
    data = request.get_json()
    path = create_word_document(
        data.get('text', ''),
        data.get('filename', 'document'),
        data.get('translated', False),
        data.get('translate_lang')
    )
    return send_file(path, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/api/tts', methods=['POST'])
def tts_endpoint():
    data = request.get_json()
    text = data.get('text', '').strip()
    if not text:
        return jsonify({'error': 'No text provided for TTS'}), 400
    path = create_audio_file(
        text,
        data.get('gender', 'male'),
        data.get('tone', 'neutral'),
        data.get('filename', 'document')
    )
    return send_file(path, as_attachment=True, mimetype='audio/mpeg')

@app.route('/health')
def health():
    return jsonify({'status': 'healthy', 'version': '5.0'})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.getenv('PORT', 5000)))
